import random
import json
import os
from docx import Document
from datetime import datetime, timedelta

class TrainingDataGenerator:
    def __init__(self):
        self.names = ['John Doe', 'Jane Smith', 'Mike Johnson', 'Sarah Williams', 'David Brown', 
                     'Emily Davis', 'Michael Wilson', 'Lisa Anderson', 'James Taylor', 'Emma Thomas']
        
        self.skills = ['Python', 'Java', 'JavaScript', 'C++', 'SQL', 'Machine Learning', 
                      'Data Analysis', 'Web Development', 'Docker', 'AWS', 'Git', 'React', 
                      'Node.js', 'TensorFlow', 'PyTorch', 'Agile', 'Scrum']
        
        self.companies = ['Tech Corp', 'Data Systems Inc.', 'Software Solutions', 'AI Innovations', 
                         'Web Dynamics', 'Cloud Technologies', 'Digital Solutions', 'Tech Startups']
        
        self.universities = ['University of Technology', 'State University', 'Technical Institute',
                           'College of Engineering', 'Institute of Science', 'National University']
        
        self.degrees = ['Bachelor of Science in Computer Science', 
                       'Master of Science in Data Science',
                       'Bachelor of Engineering in Software Engineering',
                       'Master of Business Administration',
                       'Bachelor of Science in Information Technology']
        
        self.job_titles = ['Software Engineer', 'Data Scientist', 'Full Stack Developer',
                          'Machine Learning Engineer', 'DevOps Engineer', 'Product Manager',
                          'Systems Architect', 'Technical Lead']

    def generate_resume(self) -> str:
        """Generate a sample resume."""
        name = random.choice(self.names)
        title = random.choice(self.job_titles)
        selected_skills = random.sample(self.skills, k=random.randint(5, 8))
        
        resume = f"""
{name}
{title}

CONTACT INFORMATION
Email: {name.lower().replace(' ', '.')}@email.com
Phone: (555) {random.randint(100, 999)}-{random.randint(1000, 9999)}
LinkedIn: linkedin.com/in/{name.lower().replace(' ', '-')}

SUMMARY
{random.choice(['Experienced', 'Skilled', 'Results-driven', 'Passionate'])} {title} with {random.randint(2, 8)} years of experience in {', '.join(random.sample(selected_skills, 2))}.

EDUCATION
{random.choice(self.degrees)}
{random.choice(self.universities)}
{2010 + random.randint(0, 12)}-{2014 + random.randint(0, 12)}

EXPERIENCE
{title}, {random.choice(self.companies)}
{2015 + random.randint(0, 7)} - Present
- {random.choice(['Developed', 'Implemented', 'Created', 'Designed'])} {random.choice(['systems', 'applications', 'solutions', 'platforms'])} using {random.choice(selected_skills)}
- {random.choice(['Led', 'Managed', 'Coordinated', 'Spearheaded'])} team of {random.randint(3, 8)} developers
- {random.choice(['Improved', 'Optimized', 'Enhanced', 'Streamlined'])} performance by {random.randint(20, 50)}%

{random.choice(self.job_titles)}, {random.choice(self.companies)}
{2012 + random.randint(0, 5)} - {2015 + random.randint(0, 7)}
- {random.choice(['Developed', 'Implemented', 'Created', 'Designed'])} {random.choice(['features', 'modules', 'components', 'services'])}
- {random.choice(['Collaborated', 'Worked', 'Partnered'])} with cross-functional teams
- {random.choice(['Reduced', 'Decreased', 'Minimized'])} {random.choice(['bugs', 'errors', 'issues'])} by {random.randint(30, 70)}%

SKILLS
{', '.join(selected_skills)}

PROJECTS
1. {random.choice(['AI-powered', 'Cloud-based', 'Enterprise', 'Mobile'])} {random.choice(['Application', 'Platform', 'System', 'Solution'])}
   - Used: {', '.join(random.sample(selected_skills, 3))}
2. {random.choice(['Automated', 'Intelligent', 'Scalable', 'Real-time'])} {random.choice(['Dashboard', 'Analytics', 'Framework', 'Tool'])}
   - Used: {', '.join(random.sample(selected_skills, 2))}
"""
        return resume

    def generate_non_resume(self) -> str:
        """Generate a non-resume document."""
        document_types = [
            self._generate_meeting_minutes,
            self._generate_project_proposal,
            self._generate_technical_documentation,
            self._generate_business_report
        ]
        return random.choice(document_types)()

    def _generate_meeting_minutes(self) -> str:
        """Generate sample meeting minutes."""
        date = datetime.now() - timedelta(days=random.randint(0, 30))
        attendees = random.sample(self.names, k=random.randint(3, 6))
        
        return f"""
Meeting Minutes
Date: {date.strftime('%B %d, %Y')}

Project Status Update Meeting
Location: {random.choice(['Conference Room A', 'Virtual Meeting', 'Room 202', 'Main Office'])}

Attendees:
{chr(10).join('- ' + name for name in attendees)}

Agenda:
1. Project Updates
2. Timeline Review
3. Resource Allocation
4. Next Steps

Discussion Points:
- Team discussed current project status and milestones
- Reviewed upcoming deadlines and resource requirements
- Identified potential risks and mitigation strategies

Action Items:
1. {random.choice(attendees)} to complete documentation by {(date + timedelta(days=random.randint(3, 10))).strftime('%B %d')}
2. {random.choice(attendees)} to schedule follow-up meetings with stakeholders
3. Team to provide weekly progress updates

Next Meeting: {(date + timedelta(days=7)).strftime('%B %d, %Y')}
"""

    def _generate_project_proposal(self) -> str:
        """Generate sample project proposal."""
        return f"""
Project Proposal
{datetime.now().strftime('%B %d, %Y')}

Project: {random.choice(['Digital Transformation', 'System Upgrade', 'Platform Migration', 'Infrastructure Enhancement'])}

Executive Summary:
This proposal outlines the implementation of a new {random.choice(['system', 'platform', 'solution', 'framework'])} to improve {random.choice(['efficiency', 'productivity', 'performance', 'scalability'])}.

Objectives:
1. Enhance system performance
2. Reduce operational costs
3. Improve user experience

Budget: ${random.randint(50, 500)}k
Timeline: {random.randint(3, 12)} months

Technical Requirements:
- {random.choice(self.skills)} integration
- {random.choice(self.skills)} implementation
- {random.choice(self.skills)} support

Risk Assessment:
- Technical challenges
- Resource constraints
- Timeline dependencies
"""

    def _generate_technical_documentation(self) -> str:
        """Generate sample technical documentation."""
        return f"""
Technical Documentation
Version: {random.randint(1, 3)}.{random.randint(0, 9)}

System Architecture Overview
{random.choice(['Cloud-based', 'Microservices', 'Monolithic', 'Serverless'])} Architecture

Components:
1. Frontend
   - Technology: {random.choice(['React', 'Angular', 'Vue.js'])}
   - Purpose: User interface and interaction

2. Backend
   - Technology: {random.choice(['Node.js', 'Python', 'Java'])}
   - Purpose: Business logic and data processing

3. Database
   - Technology: {random.choice(['PostgreSQL', 'MongoDB', 'MySQL'])}
   - Purpose: Data storage and retrieval

API Documentation:
GET /api/v1/resources
POST /api/v1/resources
PUT /api/v1/resources/{'{id}'}
DELETE /api/v1/resources/{'{id}'}

Deployment Instructions:
1. Clone repository
2. Install dependencies
3. Configure environment
4. Run tests
5. Deploy to production
"""

    def _generate_business_report(self) -> str:
        """Generate sample business report."""
        return f"""
Quarterly Business Report
Q{random.randint(1, 4)} {datetime.now().year}

Executive Summary:
This report presents the quarterly performance analysis and key metrics for {random.choice(self.companies)}.

Financial Highlights:
- Revenue: ${random.randint(1, 10)}M
- Growth: {random.randint(5, 25)}%
- Operating Costs: ${random.randint(500, 900)}k

Key Achievements:
1. Launched {random.randint(2, 5)} new products
2. Expanded to {random.randint(2, 5)} new markets
3. Increased customer base by {random.randint(10, 30)}%

Challenges:
- Market competition
- Resource allocation
- Technical debt

Recommendations:
1. Invest in new technologies
2. Expand team capacity
3. Improve processes
"""

def create_training_dataset(output_dir: str, num_samples: int = 50):
    """Create a training dataset with equal numbers of resumes and non-resumes."""
    os.makedirs(output_dir, exist_ok=True)
    generator = TrainingDataGenerator()
    
    # Create samples
    for i in range(num_samples):
        # Generate resume
        resume = generator.generate_resume()
        resume_doc = Document()
        resume_doc.add_paragraph(resume)
        resume_doc.save(os.path.join(output_dir, f'resume_{i+1}.docx'))
        
        # Generate non-resume
        non_resume = generator.generate_non_resume()
        non_resume_doc = Document()
        non_resume_doc.add_paragraph(non_resume)
        non_resume_doc.save(os.path.join(output_dir, f'non_resume_{i+1}.docx'))
    
    # Create labels file
    labels = {
        f'resume_{i+1}.docx': 1 for i in range(num_samples)
    }
    labels.update({
        f'non_resume_{i+1}.docx': 0 for i in range(num_samples)
    })
    
    with open(os.path.join(output_dir, 'labels.json'), 'w') as f:
        json.dump(labels, f, indent=2)
    
    print(f"Created {num_samples} resumes and {num_samples} non-resumes in {output_dir}")
    print(f"Labels saved to {os.path.join(output_dir, 'labels.json')}")

if __name__ == "__main__":
    # Create training dataset
    output_dir = "training_data"
    create_training_dataset(output_dir, num_samples=50)  # Creates 50 resumes and 50 non-resumes
